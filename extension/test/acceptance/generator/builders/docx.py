"""DOCX builders — OOXML concealment vectors from the Threat Model.

Vectors (CLAUDE.md DOCX doctrine): the Word "hidden" attribute (w:vanish),
white-on-white colour (ΔE<25), and micro-font (<2pt). Multi-run splitting models
a payload broken across runs so no single run is the whole instruction.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor

from payloads import BODY_EN


def _doc(body: list[str] = BODY_EN) -> Document:
    d = Document()
    for line in body:
        d.add_paragraph(line)
    return d


def save(doc: Document, out) -> None:
    doc.save(out)


# ── TRUE-POSITIVE vectors ─────────────────────────────────────────────────────
def vanish(payload: str, body: list[str] = BODY_EN) -> Document:
    """w:vanish — the Word 'hidden' attribute."""
    d = _doc(body)
    d.add_paragraph().add_run(payload).font.hidden = True
    return d


def white(payload: str, body: list[str] = BODY_EN) -> Document:
    """White text on the default white background."""
    d = _doc(body)
    d.add_paragraph().add_run(payload).font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return d


def tiny(payload: str, *, pt: float = 1.0) -> Document:
    """Micro-font run (<2pt)."""
    d = _doc()
    d.add_paragraph().add_run(payload).font.size = Pt(pt)
    return d


def color(payload: str, rgb: tuple[int, int, int]) -> Document:
    d = _doc()
    d.add_paragraph().add_run(payload).font.color.rgb = RGBColor(*rgb)
    return d


def split_runs(prefix: str, payload: str) -> Document:
    """Visible run, then a hidden adversarial run in the SAME paragraph."""
    d = _doc()
    p = d.add_paragraph()
    p.add_run(prefix)
    p.add_run(payload).font.hidden = True
    return d


# ── TRUE-NEGATIVE / FP-hardening ──────────────────────────────────────────────
def plain(body: list[str] = BODY_EN) -> Document:
    return _doc(body)


def visible_extra(text: str, body: list[str] = BODY_EN) -> Document:
    d = _doc(body)
    d.add_paragraph(text)
    return d


def innocent_concealed(note: str, mode: str) -> Document:
    """Concealed but innocent (no AI intent) — must stay clean via AND-gate."""
    d = _doc()
    run = d.add_paragraph().add_run(note)
    if mode == "white":
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif mode == "vanish":
        run.font.hidden = True
    elif mode == "tiny":
        run.font.size = Pt(1)
    return d


# ── Robustness ────────────────────────────────────────────────────────────────
def corrupt_bytes() -> bytes:
    return b"PK\x03\x04" + bytes(range(256)) * 4


def encrypted_bytes() -> bytes:
    # OLE2 / CFB container magic — what a password-encrypted .docx looks like.
    return bytes([0xD0, 0xCF, 0x11, 0xE0, 0xA1, 0xB1, 0x1A, 0xE1]) + bytes(64)
